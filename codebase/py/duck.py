from docx import Document
from docx.shared import Inches

TITLE = "$TITLE$"
ABSTRACT = """Jody fox
is going dpwn
the hill"""
KEYWORDS = ','.join(['a', 'b', 'c', 'd'])

document = Document()
section = document.sections[0]
header = section.header
paragraph = header.paragraphs[0]
paragraph.text = "Left Text\tCenter Text\tRight Text"
paragraph.style = document.styles["Header"]

document.add_heading(TITLE, level=0)
document.add_paragraph("Keywords: " + KEYWORDS,  style='Intense Quote')
document.add_heading("Abstract", level=1)
abstract = document.add_paragraph(ABSTRACT)
p = abstract

document.add_heading("Introduction", level=1)
document.add_heading("Method", level=1)
document.add_heading("Results", level=1)
document.add_heading("Conclusions", level=1)

p.add_run('bold').bold = True
p.add_run(' and some ')
p.add_run('italic.').italic = True

document.add_paragraph(
    'first item in unordered list', style='List Bullet'
)
document.add_paragraph(
    'first item in ordered list', style='List Number'
)

#document.add_picture('monty-truth.png', width=Inches(1.25))

records = (
    (3, '101', 'Spam'),
    (7, '422', 'Eggs'),
    (4, '631', 'Spam, spam, eggs, and spam')
)

table = document.add_table(rows=1, cols=3)
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Qty'
hdr_cells[1].text = 'Id'
hdr_cells[2].text = 'Desc'
for qty, id, desc in records:
    row_cells = table.add_row().cells
    row_cells[0].text = str(qty)
    row_cells[1].text = id
    row_cells[2].text = desc
table.style = 'Table Grid'

document.add_page_break()

document.save('demo.docx')
